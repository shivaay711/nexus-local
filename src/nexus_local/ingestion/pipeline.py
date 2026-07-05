"""Document ingestion: parse -> hash -> dedup -> chunk -> persist -> index.

PDF parsing uses a fallback chain (pdfplumber -> pypdf). Parser confidence is
recorded honestly; extraction from PDFs is never claimed to be perfect.
"""
from __future__ import annotations

import csv
import hashlib
import io
import json
from dataclasses import dataclass
from pathlib import Path

from sqlalchemy import select, text as sqltext

from ..config import settings
from ..db import audit, get_session
from ..models import Chunk, Document
from ..security.file_validation import validate_import_path


@dataclass
class ParsedPage:
    text: str
    page: int | None = None


@dataclass
class ParseResult:
    pages: list[ParsedPage]
    parser: str
    confidence: float
    error: str | None = None


def sha256_file(path: Path) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for block in iter(lambda: f.read(1 << 20), b""):
            h.update(block)
    return h.hexdigest()


# ---------------------------------------------------------------- parsers
def _parse_pdf(path: Path) -> ParseResult:
    try:
        import pdfplumber
        pages = []
        with pdfplumber.open(path) as pdf:
            for i, page in enumerate(pdf.pages, 1):
                pages.append(ParsedPage(text=page.extract_text() or "", page=i))
        if any(p.text.strip() for p in pages):
            return ParseResult(pages, "pdfplumber", 0.85)
    except Exception:
        pass
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        pages = [ParsedPage(text=p.extract_text() or "", page=i)
                 for i, p in enumerate(reader.pages, 1)]
        if any(p.text.strip() for p in pages):
            return ParseResult(pages, "pypdf", 0.7)
        return ParseResult(pages, "pypdf", 0.2,
                           error="No extractable text; PDF may be scanned (no OCR configured)")
    except Exception as e:  # both parsers failed
        return ParseResult([], "none", 0.0, error=f"PDF parse failed: {e}")


def _parse_docx(path: Path) -> ParseResult:
    try:
        import docx
        d = docx.Document(str(path))
        body = "\n".join(p.text for p in d.paragraphs)
        return ParseResult([ParsedPage(body)], "python-docx", 0.9)
    except Exception as e:
        return ParseResult([], "python-docx", 0.0, error=str(e))


def _parse_text(path: Path) -> ParseResult:
    body = path.read_text(encoding="utf-8", errors="replace")
    return ParseResult([ParsedPage(body)], "plaintext", 1.0)


def _parse_csv(path: Path) -> ParseResult:
    rows = list(csv.reader(io.StringIO(path.read_text(encoding="utf-8", errors="replace"))))
    if not rows:
        return ParseResult([], "csv", 0.0, error="Empty CSV")
    header = rows[0]
    lines = [", ".join(header)]
    for row in rows[1:]:
        lines.append("; ".join(f"{h}={v}" for h, v in zip(header, row)))
    return ParseResult([ParsedPage("\n".join(lines))], "csv", 0.95)


def _parse_json(path: Path) -> ParseResult:
    try:
        data = json.loads(path.read_text(encoding="utf-8", errors="replace"))
        return ParseResult([ParsedPage(json.dumps(data, indent=2)[:200_000])], "json", 0.95)
    except json.JSONDecodeError as e:
        return ParseResult([ParsedPage(path.read_text(errors="replace"))], "json-as-text", 0.5,
                           error=f"Invalid JSON, indexed as text: {e}")


def _parse_notebook(path: Path) -> ParseResult:
    try:
        nb = json.loads(path.read_text(encoding="utf-8", errors="replace"))
        parts = []
        for cell in nb.get("cells", []):
            src = "".join(cell.get("source", []))
            tag = "```python\n{}\n```" if cell.get("cell_type") == "code" else "{}"
            parts.append(tag.format(src))
        return ParseResult([ParsedPage("\n\n".join(parts))], "ipynb", 0.9)
    except Exception as e:
        return ParseResult([], "ipynb", 0.0, error=str(e))


def _parse_html(path: Path) -> ParseResult:
    import re
    raw = path.read_text(encoding="utf-8", errors="replace")
    raw = re.sub(r"<(script|style)[^>]*>.*?</\1>", " ", raw, flags=re.S | re.I)
    txt = re.sub(r"<[^>]+>", " ", raw)
    txt = re.sub(r"\s+", " ", txt).strip()
    return ParseResult([ParsedPage(txt)], "html-strip", 0.7)


PARSERS = {
    ".pdf": _parse_pdf, ".docx": _parse_docx, ".txt": _parse_text, ".md": _parse_text,
    ".py": _parse_text, ".js": _parse_text, ".ts": _parse_text, ".csv": _parse_csv,
    ".json": _parse_json, ".jsonl": _parse_text, ".ipynb": _parse_notebook, ".html": _parse_html,
}


# ---------------------------------------------------------------- chunking
def chunk_text(body: str, size: int | None = None, overlap: int | None = None) -> list[str]:
    """Paragraph-aware sliding window chunker with character budgets."""
    size = size or settings.chunk_size_chars
    overlap = overlap or settings.chunk_overlap_chars
    paragraphs = [p.strip() for p in body.split("\n\n") if p.strip()]
    chunks: list[str] = []
    buf = ""
    for p in paragraphs:
        while len(p) > size:  # split any monster paragraph
            head, p = p[:size], p[max(0, size - overlap):]
            if buf:
                chunks.append(buf)
                buf = ""
            chunks.append(head)
        if len(buf) + len(p) + 2 <= size:
            buf = f"{buf}\n\n{p}" if buf else p
        else:
            chunks.append(buf)
            tail = buf[-overlap:] if overlap and buf else ""
            buf = (tail + "\n\n" + p).strip()
            if len(buf) > size:
                buf = p
    if buf:
        chunks.append(buf)
    return [c for c in chunks if c.strip()]


def _heading_for(chunk: str) -> str | None:
    for line in chunk.splitlines():
        line = line.strip()
        if line.startswith("#"):
            return line.lstrip("#").strip()[:200]
    return None


# ---------------------------------------------------------------- pipeline
def import_document(path: Path | str) -> Document:
    """Validate, parse, dedup by SHA-256, chunk, persist, and lexically index.
    Embeddings are added by the retrieval layer (see retrieval.index)."""
    p = validate_import_path(Path(path))
    digest = sha256_file(p)
    with get_session() as s:
        existing = s.scalar(select(Document).where(
            Document.sha256 == digest, Document.deleted.is_(False)))
        if existing:
            return existing  # idempotent import

        parser = PARSERS.get(p.suffix.lower())
        if parser is None:
            raise ValueError(f"No parser for {p.suffix}")
        result = parser(p)

        doc = Document(
            path=str(p), filename=p.name, sha256=digest, size_bytes=p.stat().st_size,
            media_type=p.suffix.lower().lstrip("."), parser_used=result.parser,
            parser_confidence=result.confidence, parse_error=result.error,
        )
        s.add(doc)
        s.flush()

        ordinal = 0
        for page in result.pages:
            for piece in chunk_text(page.text):
                ch = Chunk(document_id=doc.id, ordinal=ordinal, text=piece,
                           page=page.page, heading=_heading_for(piece))
                s.add(ch)
                s.flush()
                s.execute(sqltext(
                    "INSERT INTO chunk_fts (chunk_id, document_id, body) VALUES (:c, :d, :b)"),
                    {"c": ch.id, "d": doc.id, "b": piece})
                ordinal += 1
        audit(s, "import", "document", doc.id, f"{p.name} chunks={ordinal} parser={result.parser}")
        s.commit()
        return doc


def delete_document(document_id: str) -> None:
    with get_session() as s:
        doc = s.get(Document, document_id)
        if not doc:
            return
        doc.deleted = True
        s.execute(sqltext("DELETE FROM chunk_fts WHERE document_id = :d"), {"d": document_id})
        audit(s, "delete", "document", document_id)
        s.commit()
